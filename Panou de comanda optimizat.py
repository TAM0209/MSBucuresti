import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from openpyxl import load_workbook
import os
from datetime import datetime
import sqlite3
import openpyxl
import pandas as pd

def search_database():
    # Funcția care este apelată când butonul "Baza de date" este apăsat
    
    # Conectarea la baza de date
    conn = sqlite3.connect('C:/Users/tomaa/OneDrive/Desktop/BazaDeDateV1/Macro/SQLite+VSC+PYTHON/GenerareCVT/Bazadedate.db')  # Schimbă 'nume_baza_de_date.db' cu numele bazei de date
    cursor = conn.cursor()
    
    # Crearea ferestrei de căutare
    search_window = tk.Toplevel(root)
    search_window.title("Căutare în baza de date")
    
    # Eticheta pentru introdus seria
    label = tk.Label(search_window, text="Introduceți seria:")
    label.pack()
    
    # Caseta de text pentru introducerea seriei
    serial_entry = tk.Entry(search_window)
    serial_entry.pack()
    
    # Funcția pentru căutare
    def search():
        serial = serial_entry.get()
        if serial:
            # Executarea interogării în baza de date
            query = f"SELECT * FROM bazadatemetron WHERE Serial = ?"
            cursor.execute(query, (serial,))
            results = cursor.fetchall()
            
            if results:
                # Salvarea rezultatelor într-un fișier Excel
                df = pd.DataFrame(results, columns=[description[0] for description in cursor.description])
                file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
                if file_path:
                    df.to_excel(file_path, index=False)
            else:
                tk.messagebox.showinfo("Informație", "Nu s-au găsit rezultate pentru seria introdusă.")
        else:
            tk.messagebox.showwarning("Avertizare", "Vă rugăm să introduceți o serie.")
    
    # Butonul pentru căutare
    search_button = tk.Button(search_window, text="Caută", command=search)
    search_button.pack()

# Funcția pentru importul datelor în tabela "procentaj"
def import_excel_in_procentaj(file_path):
    
    # Conectarea la baza de date
    conn = sqlite3.connect('C:/Users/tomaa/OneDrive/Desktop/BazaDeDateV1/Macro/SQLite+VSC+PYTHON/GenerareCVT/Bazadedate.db')
    cursor = conn.cursor()

    # Citirea datelor din fișierul Excel
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    # Iterăm prin fiecare rând al fișierului Excel, începând cu al doilea rând
    for row_num in range(2, sheet.max_row + 1):
        # Citim valorile din fiecare coloană a rândului curent din fișierul Excel
        values = [sheet.cell(row=row_num, column=col_num).value for col_num in range(1, sheet.max_column + 1)]

        # Inserăm valorile citite în tabela "procentaj" din baza de date
        cursor.execute("INSERT INTO procentaj VALUES (?, ?, ?, ?)", values)

    # Salvăm schimbările și închidem conexiunea la baza de date
    conn.commit()
    conn.close()

# Funcția care este apelată când butonul "Import procentaj AT" este apăsat
def on_import_procentaj_AT_button_click():
    # Deschidem o fereastră pentru selectarea fișierului Excel
    file_path = filedialog.askopenfilename(filetypes=[("Fișiere Excel", "*.xlsx"), ("Toate fișierele", "*.*")])

    # Dacă utilizatorul a selectat un fișier, îl importăm în tabela "Import procentaj AT"
    if file_path:
        import_excel_in_procentaj(file_path)
        # Afisăm un mesaj de confirmare
        messagebox.showinfo("Salvare completă", "Datele au fost salvate cu succes în baza de date!")

# Funcția pentru importul datelor în tabela "mize"
def import_excel_in_mize(file_path):
    
    # Conectarea la baza de date
    conn = sqlite3.connect('C:/Users/tomaa/OneDrive/Desktop/BazaDeDateV1/Macro/SQLite+VSC+PYTHON/GenerareCVT/Bazadedate.db')
    cursor = conn.cursor()

    # Citirea datelor din fișierul Excel
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    # Iterăm prin fiecare rând al fișierului Excel, începând cu al doilea rând
    for row_num in range(2, sheet.max_row + 1):
        # Citim valorile din fiecare coloană a rândului curent din fișierul Excel
        values = [sheet.cell(row=row_num, column=col_num).value for col_num in range(1, sheet.max_column + 1)]

        # Inserăm valorile citite în tabela "mize" din baza de date
        cursor.execute("INSERT INTO mize VALUES (?, ?, ?, ?, ?, ?)", values)

    # Salvăm schimbările și închidem conexiunea la baza de date
    conn.commit()
    conn.close()

# Funcția care este apelată când butonul "Import mize AT" este apăsat
def on_import_mize_AT_button_click():
    # Deschidem o fereastră pentru selectarea fișierului Excel
    file_path = filedialog.askopenfilename(filetypes=[("Fișiere Excel", "*.xlsx"), ("Toate fișierele", "*.*")])

    # Dacă utilizatorul a selectat un fișier, îl importăm în tabela "Import mize AT"
    if file_path:
        import_excel_in_mize(file_path)
        # Afisăm un mesaj de confirmare
        messagebox.showinfo("Salvare completă", "Datele au fost salvate cu succes în baza de date!")

# Funcția pentru importul datelor în tabela "AT"
def import_excel_in_AT(file_path):
    
    # Conectarea la baza de date
    conn = sqlite3.connect('C:/Users/tomaa/OneDrive/Desktop/BazaDeDateV1/Macro/SQLite+VSC+PYTHON/GenerareCVT/Bazadedate.db')
    cursor = conn.cursor()

    # Citirea datelor din fișierul Excel
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    # Iterăm prin fiecare rând al fișierului Excel, începând cu al doilea rând
    for row_num in range(2, sheet.max_row + 1):
        # Citim valorile din fiecare coloană a rândului curent din fișierul Excel
        values = [sheet.cell(row=row_num, column=col_num).value for col_num in range(1, sheet.max_column + 1)]

        # Inserăm valorile citite în tabela "AT" din baza de date
        cursor.execute("INSERT INTO AT VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", values)

    # Salvăm schimbările și închidem conexiunea la baza de date
    conn.commit()
    conn.close()

# Funcția care este apelată când butonul "Iformatii AT" este apăsat
def on_informatii_AT_button_click():
    # Deschidem o fereastră pentru selectarea fișierului Excel
    file_path = filedialog.askopenfilename(filetypes=[("Fișiere Excel", "*.xlsx"), ("Toate fișierele", "*.*")])

    # Dacă utilizatorul a selectat un fișier, îl importăm în tabela "AT"
    if file_path:
        import_excel_in_AT(file_path)
        # Afisăm un mesaj de confirmare
        messagebox.showinfo("Salvare completă", "Datele au fost salvate cu succes în baza de date!")

# Fereastră de dialog personalizată pentru introducerea datelor
class ProprietariDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Introducere Date Proprietari")
        self.geometry("300x200")
        
        # Etichete și câmpuri de introducere pentru datele proprietarilor
        tk.Label(self, text="Companie:").grid(row=0, column=0, sticky="w")
        self.companie_entry = tk.Entry(self)
        self.companie_entry.grid(row=0, column=1)
        
        tk.Label(self, text="Adresa:").grid(row=1, column=0, sticky="w")
        self.adresa_entry = tk.Entry(self)
        self.adresa_entry.grid(row=1, column=1)
        
        tk.Label(self, text="CUI:").grid(row=2, column=0, sticky="w")
        self.cui_entry = tk.Entry(self)
        self.cui_entry.grid(row=2, column=1)
        
        tk.Label(self, text="Licență:").grid(row=3, column=0, sticky="w")
        self.licenta_entry = tk.Entry(self)
        self.licenta_entry.grid(row=3, column=1)
        
        # Butonul de salvare
        save_button = tk.Button(self, text="Save", command=self.save_data)
        save_button.grid(row=4, columnspan=2, pady=10)
        
    # Funcție pentru salvarea datelor în baza de date
    def save_data(self):
        companie = self.companie_entry.get()
        adresa = self.adresa_entry.get()
        cui = self.cui_entry.get()
        licenta = self.licenta_entry.get()
        
        # Verificăm dacă toate câmpurile sunt completate
        if companie and adresa and cui and licenta:
            # Conexiune la baza de date SQLite
            conn = sqlite3.connect('C:/Users/tomaa/OneDrive/Desktop/BazaDeDateV1/Macro/SQLite+VSC+PYTHON/GenerareCVT/Bazadedate.db')
            cursor = conn.cursor()
            
            # Inserare în baza de date
            cursor.execute("INSERT INTO 'Proprietari' ('Proprietari', 'AdresaProrpietari', 'CUIProprietari', 'LicentaProprietari') VALUES (?, ?, ?, ?)", (companie, adresa, cui, licenta))
            cursor.execute("INSERT INTO 'organizatori' ('Organizatori', 'AdresaOrganizatori', 'CUIOrganizatori', 'LicentaOrganizatori') VALUES (?, ?, ?, ?)", (companie, adresa, cui, licenta))

            # Salvăm schimbările și închidem conexiunea
            conn.commit()
            conn.close()
            
            # Închidem fereastra de dialog
            self.destroy()
            
            # Afisăm un mesaj de confirmare
            messagebox.showinfo("Salvare completă", "Datele au fost salvate cu succes în baza de date!")
        else:
            # Afisăm un mesaj de eroare dacă nu toate câmpurile sunt completate
            messagebox.showerror("Eroare", "Vă rugăm să completați toate câmpurile!")


def sterge_date_excel_dupa_import(fisier_excel):
    # Încarcăm fișierul Excel
    workbook = openpyxl.load_workbook(fisier_excel)
    sheet = workbook.active
    
    # Iterăm peste fiecare rând începând cu al doilea și fiecare coloană începând cu a doua
    for row in sheet.iter_rows(min_row=2, min_col=1):
        for cell in row:
            # Setăm valoarea celulei la None
            cell.value = None
    
    # Salvăm modificările
    workbook.save(fisier_excel)

def verify_data_in_database(sheet):
    # Deschide conexiunea la baza de date
    conn = sqlite3.connect(os.path.join(folder_curent, 'Bazadedate.db'))
    cursor = conn.cursor()

    # Verificare informații din coloana 4 (AT)
    for row_num in range(2, sheet.max_row + 1):
        valoare_coloana_4 = sheet.cell(row=row_num, column=4).value
        if valoare_coloana_4:
            cursor.execute("SELECT * FROM AT WHERE AT = ?", (valoare_coloana_4,))
            rezultate = cursor.fetchall()
            if not rezultate:
                messagebox.showwarning("Avertisment", f"Informația din coloana 4 (AT) pentru rândul {row_num} lipsește din baza de date. Valoare din fișierul Excel: {valoare_coloana_4}")

    # Verificare informații din coloana 17 (Proprietari)
    for row_num in range(2, sheet.max_row + 1):
        valoare_coloana_17 = sheet.cell(row=row_num, column=17).value
        if valoare_coloana_17:
            cursor.execute("SELECT * FROM Proprietari WHERE Proprietari = ?", (valoare_coloana_17,))
            rezultate = cursor.fetchall()
            if not rezultate:
                messagebox.showwarning("Avertisment", f"Informația din coloana 17 (Proprietari) pentru rândul {row_num} lipsește din baza de date. Valoare din fișierul Excel: {valoare_coloana_17}")

    # Verificare informații din coloana 18 (Organizatori)
    for row_num in range(2, sheet.max_row + 1):
        valoare_coloana_18 = sheet.cell(row=row_num, column=18).value
        if valoare_coloana_18:
            cursor.execute("SELECT * FROM organizatori WHERE Organizatori = ?", (valoare_coloana_18,))
            rezultate = cursor.fetchall()
            if not rezultate:
                messagebox.showwarning("Avertisment", f"Informația din coloana 18 (Organizatori) pentru rândul {row_num} lipsește din baza de date. Valoare din fișierul Excel: {valoare_coloana_18}")

    # Închide conexiunea la baza de date
    conn.close()

# Funcția pentru importul datelor din Excel în fișierul BOOM.xlsx
def import_excel_in_BOOM(cale_fisier_importat, cale_fisier_BOOM):
    # Încărcăm fișierul Excel importat
    workbook_importat = load_workbook(cale_fisier_importat)
    sheet_importat = workbook_importat.active

    # Determinăm numărul total de rânduri în fișierul importat
    numar_rânduri_importate = sheet_importat.max_row

    # Determinăm numărul total de coloane în fișierul importat
    numar_coloane_importate = sheet_importat.max_column

    # Încărcăm fișierul Excel destinatar (BOOM.xlsx)
    workbook_destinatar = load_workbook(cale_fisier_BOOM)
    sheet_destinatar = workbook_destinatar.active

    # Determinăm numărul total de rânduri în fișierul destinatar
    numar_rânduri_destinatar = sheet_destinatar.max_row

    # Copiem datele din fișierul importat în fișierul destinatar (BOOM.xlsx)
    for i in range(2, numar_rânduri_importate + 1):
        for j in range(1, numar_coloane_importate + 1):
            valoare_celulă = sheet_importat.cell(row=i, column=j).value
            # Verificăm dacă valoarea este o dată și o formatăm dacă este necesar
            if isinstance(valoare_celulă, datetime):
                valoare_celulă = valoare_celulă.strftime("%d.%m.%Y")
            sheet_destinatar.cell(row=numar_rânduri_destinatar + i - 1, column=j).value = valoare_celulă

    # Salvăm schimbările în fișierul destinatar (BOOM.xlsx)
    workbook_destinatar.save(cale_fisier_BOOM)
    # Afisam un mesaj de confirmare
    messagebox.showinfo("Import complet", "Datele au fost importate cu succes în BOOM.xlsx!")

    # După ce datele sunt importate, se pot face verificările
    verify_data_in_database(sheet_importat)  # Schimbare aici

# Funcția pentru generarea unor date
def generate_data():
    # Aici poți adăuga funcționalitatea specifică pentru generarea datelor
    from openpyxl import load_workbook
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import os
    import sqlite3
    import datetime

    # Obținem calea către directorul curent
    folder_curent = os.path.dirname(os.path.abspath(__file__))

    # Conexiunea la baza de date SQLite
    conn = sqlite3.connect(os.path.join(folder_curent, 'Bazadedate.db'))
    cursor = conn.cursor()

    # Verificăm tabelele din baza de date
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tabele = cursor.fetchall()

    # Încărcăm fișierul Excel
    workbook = load_workbook(os.path.join(folder_curent, 'BOOM.xlsx'))
    sheet = workbook['Sheet1']

    # Iterăm prin fiecare rând din fișierul Excel și inserăm datele în baza de date
    for row in sheet.iter_rows(min_row=2, values_only=True):  # Excludem rândul de antet
        # Inserăm datele în tabela "bazadatemetron"
        cursor.execute("""
            INSERT INTO bazadatemetron 
            (serial, Year, producator, AT, marca_cvt, marca_talon, marca_memorie, marca_cpu, marca_contori, marca_placuta, 
            marca_eticheta_conformitate, metrolog, serie_marca_autentificare, tip_verificare, data_verificarii, data_generarii, 
            proprietar, organizator, doc_prov, adresa, mod_achizitie, separat_impreuna_cu, soft_produs_de, achizitionat_de_la, 
            denominatii, denominatie_testare, bancnota_min, bancnota_max, tip_creditare, valoare_jucata, multiplicator, contor_in, 
            contor_out, contor_remote, Ram_Clear, CMD) 
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, row)

    # Încărcăm macheta Word
    cale_macheta_word = os.path.join(folder_curent, 'CVT404.docx')

    # Iterăm prin fiecare rând din Excel începând cu al doilea rând
    for row_num in range(2, sheet.max_row + 1):
        # Creăm un document Word nou pentru fiecare rând din Excel
        document_macheta = Document(cale_macheta_word)
        
        # Căutăm și înlocuim marcajele din întregul document Word, inclusiv în tabele
        for paragraph in document_macheta.paragraphs:
            for col_num in range(1, sheet.max_column + 1):
                placeholder = "{" + sheet.cell(row=1, column=col_num).value + "}"
                if placeholder in paragraph.text:
                    cell_value = sheet.cell(row=row_num, column=col_num).value
                    # Verificăm dacă valoarea din celulă este None sau nu
                    if cell_value is not None:
                        if isinstance(cell_value, datetime.datetime):
                            # Formatarea datei în formatul "DD.MM.YYYY"
                            formatted_date = cell_value.strftime("%d.%m.%Y")
                            paragraph.text = paragraph.text.replace(placeholder, formatted_date)
                        else:
                            paragraph.text = paragraph.text.replace(placeholder, str(cell_value))
                    else:
                        # Dacă valoarea este None, lăsăm marcajul necompletat
                        paragraph.text = paragraph.text.replace(placeholder, "")
        
        # Căutăm și înlocuim marcajele în tabelele documentului Word
        for table in document_macheta.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Iterăm prin fiecare paragraf din celula tabelului
                    for paragraph in cell.paragraphs:
                        for col_num in range(1, sheet.max_column + 1):
                            placeholder = "{" + sheet.cell(row=1, column=col_num).value + "}"
                            if placeholder in paragraph.text:
                                cell_value = sheet.cell(row=row_num, column=col_num).value
                                # Verificăm dacă valoarea din celulă este None sau nu
                                if cell_value is not None:
                                    paragraph.text = paragraph.text.replace(placeholder, str(cell_value))
                                else:
                                    # Dacă valoarea este None, lăsăm marcajul necompletat
                                    paragraph.text = paragraph.text.replace(placeholder, "")

    # Calculăm rezultatul formulei și îl introducem în locul marcajului "{CreditPoints}"
        coloana26 = float(sheet.cell(row=row_num, column=26).value)    
        coloana30 = float(sheet.cell(row=row_num, column=30).value)
        rezultat_formula = coloana30 / coloana26

        # Asigurăm că rezultatul final
        rezultat_final = rezultat_formula

        # Căutăm marcajul "{CreditPoints}" în tabelele documentului Word și îl înlocuim cu rezultatul formulei
        for table in document_macheta.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{CreditPoints}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{CreditPoints}", str(int(rezultat_final)))

    # Calculăm rezultatul formulei și îl introducem în locul marcajului "{cif}"
        coloana32 = float(sheet.cell(row=row_num, column=32).value)
        coloana30 = float(sheet.cell(row=row_num, column=30).value)
        coloana31 = float(sheet.cell(row=row_num, column=31).value)
        rezultat_formula = coloana32 + coloana30 / (coloana31 / 100)
        
        # Asigurăm că rezultatul are întotdeauna 7 cifre și adăugăm zerouri în față dacă este necesar
        rezultat_formatat = "{:.0f}".format(rezultat_formula)
        rezultat_final = rezultat_formatat.zfill(7)
        
        # Căutăm marcajul "{cif}" în tabelele documentului Word și îl înlocuim cu rezultatul formulei
        for table in document_macheta.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{cif}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{cif}", str(rezultat_final))

        # Calculăm rezultatul formulei și îl introducem în locul marcajului "{cof}"
        coloana33 = float(sheet.cell(row=row_num, column=33).value)
        coloana30 = float(sheet.cell(row=row_num, column=30).value)
        coloana31 = float(sheet.cell(row=row_num, column=31).value)
        rezultat_formula = coloana33 + coloana30 / (coloana31 / 100)
        
        # Asigurăm că rezultatul are întotdeauna 7 cifre și adăugăm zerouri în față dacă este necesar
        rezultat_formatat = "{:.0f}".format(rezultat_formula)
        rezultat_final = rezultat_formatat.zfill(7)
        
        # Căutăm marcajul "{cof}" în tabelele documentului Word și îl înlocuim cu rezultatul formulei
        for table in document_macheta.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{cof}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{cof}", str(rezultat_final))

        # Căutăm și înlocuim marcajele din toate celulele din tabelul din documentul Word
        for table in document_macheta.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Iterăm prin fiecare paragraf din celula tabelului
                    for paragraph in cell.paragraphs:
                        # Căutăm și înlocuim marcajele
                        for col_num in range(1, sheet.max_column + 1):
                            placeholder = "{" + sheet.cell(row=1, column=col_num).value + "}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(sheet.cell(row=row_num, column=col_num).value))

    # Adăugăm datele din baza de date din tabela "mize" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=4).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "mize" valoarea din prima coloană care să corespundă cu valoarea din a patra coloană a Excel-ului
            query_mize = "SELECT [Nr_Crt] FROM mize WHERE [AT] = ? ORDER BY [Nr_Crt] DESC LIMIT 1"
            cursor.execute(query_mize, (valoare_coloana_excel,))
            ultimul_raspuns = cursor.fetchone()
            
            
            if ultimul_raspuns:
                valoare_din_a_doua_coloana = ultimul_raspuns[0]  # A doua coloană din rândul găsit
                
                
                # Găsim marcajul "{NumarSub}" în documentul Word și înlocuim-l cu valoarea din a doua coloană
                for table in document_macheta.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if "{NumarSub}" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("{NumarSub}", str(valoare_din_a_doua_coloana))

        # Adăugăm datele din baza de date din tabela "mize" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=4).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "mize" folosind valoarea din a patra coloană a Excel-ului
            query_mize = f"SELECT * FROM mize WHERE AT = ?"
            cursor.execute(query_mize, (valoare_coloana_excel,))
            rezultate_mize = cursor.fetchall()
        
        
            if rezultate_mize:
                # Găsim tabelul din document
                for table in document_macheta.tables:
                    if table.cell(0, 0).text.strip() == 'Nr. Crt':
                        # Iterăm prin rezultatele interogării bazei de date "mize"
                        for entry in rezultate_mize:
                            # Adăugăm un rând nou în tabel
                            row_cells = table.add_row().cells
                            # Populăm celulele rândului cu valorile corespunzătoare din baza de date "mize"
                            row_cells[0].text = str(entry[1])  # Nr_Crt
                            row_cells[1].text = entry[2]  # Subrograme
                            row_cells[2].text = entry[3]  # Miza_minima
                            row_cells[3].text = entry[4]  # Miza_maxima
                            row_cells[4].text = entry[5]  # Castig_maxim

                            # Centrăm textul în fiecare celulă
                            for cell in row_cells:
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adăugăm datele din baza de date din tabela "procentaj" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=4).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "procentaj" folosind valoarea din a patra coloană a Excel-ului
            query_procentaj = f"SELECT * FROM procentaj WHERE AT = ?"
            cursor.execute(query_procentaj, (valoare_coloana_excel,))
            rezultate_procentaj = cursor.fetchall()
        
        
            if rezultate_procentaj:
                # Găsim tabelul din document
                for table in document_macheta.tables:
                    if table.cell(0, 0).text.strip() == 'NUME SUBPROGRAM':
                        # Iterăm prin rezultatele interogării bazei de date "procentaj"
                        for entry in rezultate_procentaj:
                            # Adăugăm un rând nou în tabel
                            row_cells = table.add_row().cells
                            # Populăm celulele rândului cu valorile corespunzătoare din baza de date "mize"
                            row_cells[0].text = str(entry[1])  # NUME_SUBPROGRAM
                            row_cells[1].text = entry[2]  # ID_tabel_plati
                            row_cells[2].text = entry[3]  # RTP%
                        
                            # Centrăm textul în fiecare celulă
                            for cell in row_cells:
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Adăugăm datele din baza de date din tabela "AT" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=4).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "AT" folosind valoarea din a patra coloană a Excel-ului
            query_at = "SELECT * FROM AT WHERE AT = ?"
            cursor.execute(query_at, (valoare_coloana_excel,))
            rezultate_at = cursor.fetchall()

            if rezultate_at:
                # Căutăm și înlocuim marcajele în paragrafele documentului Word
                for paragraph in document_macheta.paragraphs:
                    for entry_at in rezultate_at:
                        for col_index, col_value in enumerate(entry_at):
                            # Ignorăm prima coloană, care este valoarea din coloana AT (prima coloană din tabela "AT")
                            if col_index > 0:
                                # Găsim marcajul corespunzător în paragraf și îl înlocuim cu valoarea din tabela "AT"
                                placeholder = "{Informatie" + chr(ord('A') + col_index - 1) + "}"
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(col_value))

                # Căutăm și înlocuim marcajele în tabelele documentului Word
                for table in document_macheta.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Iterăm prin fiecare paragraf din celula tabelului
                            for paragraph in cell.paragraphs:
                                for entry_at in rezultate_at:
                                    for col_index, col_value in enumerate(entry_at):
                                        # Ignorăm prima coloană, care este valoarea din coloana AT (prima coloană din tabela "AT")
                                        if col_index > 0:
                                            # Găsim marcajul corespunzător în paragraf și îl înlocuim cu valoarea din tabela "AT"
                                            placeholder = "{Informatie" + chr(ord('A') + col_index - 1) + "}"
                                            if placeholder in paragraph.text:
                                                paragraph.text = paragraph.text.replace(placeholder, str(col_value))
        
        # Adăugăm datele din baza de date din tabela "Proprietari" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=17).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "Proprietari" folosind valoarea din a 17-a coloană a Excel-ului
            query_proprietari = "SELECT * FROM Proprietari WHERE Proprietari = ?"  # Înlocuiește "ID" cu numele real al primei coloane din tabela "Proprietari"
            cursor.execute(query_proprietari, (valoare_coloana_excel,))
            rezultate_proprietari = cursor.fetchall()

            if rezultate_proprietari:
                # Căutăm și înlocuim marcajele în paragrafele documentului Word
                for paragraph in document_macheta.paragraphs:
                    for entry_proprietar in rezultate_proprietari:
                        # Înlocuim marcajele cu valorile corespunzătoare
                        if "{AdresaProprietari}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{AdresaProprietari}", str(entry_proprietar[1]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"
                        if "{CUIProprietari}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{CUIProprietari}", str(entry_proprietar[2]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"
                        if "{LicentaProprietari}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{LicentaProprietari}", str(entry_proprietar[3]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"

                # Căutăm și înlocuim marcajele în tabelele documentului Word
                for table in document_macheta.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Iterăm prin fiecare paragraf din celula tabelului
                            for paragraph in cell.paragraphs:
                                for entry_proprietar in rezultate_proprietari:
                                    # Înlocuim marcajele cu valorile corespunzătoare
                                    if "{AdresaProprietari}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{AdresaProprietari}", str(entry_proprietar[1]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"
                                    if "{CUIProprietari}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{CUIProprietari}", str(entry_proprietar[2]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"
                                    if "{LicentaProprietari}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{LicentaProprietari}", str(entry_proprietar[3]))  # Înlocuiește indexul cu coloana reală din tabela "Proprietari"

    # Adăugăm datele din baza de date din tabela "organizatori" în documentul Word
        valoare_coloana_excel = sheet.cell(row=row_num, column=18).value
        if valoare_coloana_excel:
            # Căutăm în baza de date "organizatori" folosind valoarea din a 18-a coloană a Excel-ului
            query_organizatori = "SELECT * FROM organizatori WHERE Organizatori = ?"  # Înlocuiește "ID" cu numele real al primei coloane din tabela "organizatori"
            cursor.execute(query_organizatori, (valoare_coloana_excel,))
            rezultate_organizatori = cursor.fetchall()

            if rezultate_organizatori:
                # Căutăm și înlocuim marcajele în paragrafele documentului Word
                for paragraph in document_macheta.paragraphs:
                    for entry_organizatori in rezultate_organizatori:
                        # Înlocuim marcajele cu valorile corespunzătoare
                        if "{AdresaOrganizatori}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{AdresaOrganizatori}", str(entry_organizatori[1]))  # Înlocuiește indexul cu coloana reală din tabela "organizatori"
                        if "{CUIOrganizatori}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{CUIOrganizatori}", str(entry_organizatori[2]))  # Înlocuiește indexul cu coloana reală din tabela "organizatori"
                        if "{LicentaOrganizatori}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{LicentaOrganizatori}", str(entry_organizatori[3]))  # Înlocuiește indexul cu coloana reală din tabela "organizatori"

                # Căutăm și înlocuim marcajele în tabelele documentului Word
                for table in document_macheta.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Iterăm prin fiecare paragraf din celula tabelului
                            for paragraph in cell.paragraphs:
                                for entry_organizatori in rezultate_organizatori:
                                    # Înlocuim marcajele cu valorile corespunzătoare
                                    if "{AdresaOrganizatori}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{AdresaOrganizatori}", str(entry_organizatori[1]))  # Înlocuiește indexul cu coloana reală din tabela "Organizatori"
                                    if "{CUIOrganizatori}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{CUIOrganizatori}", str(entry_organizatori[2]))  # Înlocuiește indexul cu coloana reală din tabela "Organizatori"
                                    if "{LicentaOrganizatori}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace("{LicentaOrganizatori}", str(entry_organizatori[3]))  # Înlocuiește indexul cu coloana reală din tabela "Organizatori"

        # Salvăm documentul Word actualizat în același folder cu scriptul
        nume_document = f"{sheet.cell(row=row_num, column=1).value}-{sheet.cell(row=row_num, column=2).value}.docx"
        cale_document = os.path.join(folder_curent, nume_document)
        document_macheta.save(cale_document)
        
    # Salvăm schimbările și închidem conexiunea la baza de date
    conn.commit()
    conn.close()
    # Ștergem datele din fișierul BOOM.xlsx după ce documentele au fost generate
    sterge_date_excel_dupa_import(os.path.join(folder_curent, 'BOOM.xlsx'))
pass

# Funcția care este apelată când butonul de import este apăsat
def on_import_button_click():
    # Deschidem o fereastră pentru selectarea fișierului Excel
    fisier_selectat = filedialog.askopenfilename(filetypes=[("Fișiere Excel", "*.xlsx"), ("Toate fișierele", "*.*")])

    # Dacă utilizatorul a selectat un fișier, îl importăm în BOOM.xlsx
    if fisier_selectat:
        import_excel_in_BOOM(fisier_selectat, os.path.join(folder_curent, 'BOOM.xlsx'))

# Funcția care este apelată când butonul de generare este apăsat
def on_generate_button_click():
    generate_data()

# Obținem calea către directorul curent
folder_curent = os.path.dirname(os.path.abspath(__file__))

# Creăm o fereastră principală tkinter
root = tk.Tk()
root.title("Aplicație de import Excel")
# Mărirea dimensiunii ferestrei principale
root.geometry("600x400")  # Schimbă dimensiunea la 600x400 pixeli sau alte dimensiuni dorite

# Adăugăm un buton pentru import
import_button = tk.Button(root, text="Import Excel", command=on_import_button_click)
import_button.pack(pady=10)

# Adăugăm un buton pentru generare
generate_button = tk.Button(root, text="Generare", command=on_generate_button_click)
generate_button.pack(pady=10)

# Adăugăm un buton pentru "Informatii AT" în interfața grafică
informatii_AT_button = tk.Button(root, text="Informatii AT", command=on_informatii_AT_button_click)
informatii_AT_button.pack(pady=10)

# Adăugăm un buton pentru "Import mize AT" în interfața grafică
import_mize_AT_button = tk.Button(root, text="Import mize AT", command=on_import_mize_AT_button_click)
import_mize_AT_button.pack(pady=10)

# Adăugăm un buton pentru "Import procentaj AT" în interfața grafică
import_procentaj_AT_button = tk.Button(root, text="Import procentaj AT", command=on_import_procentaj_AT_button_click)
import_procentaj_AT_button.pack(pady=10)

# Funcția care este apelată când butonul "Proprietari/Organizatori" este apăsat
def on_proprietari_button_click():
    # Afisăm fereastra de dialog personalizată pentru introducerea datelor
    dialog = ProprietariDialog(root)

proprietari_button = tk.Button(root, text="Proprietari/Organizatori", command=on_proprietari_button_click)
proprietari_button.pack(pady=20)

# Butonul pentru căutare în bază de date
search_button = tk.Button(root, text="Baza de date", command=search_database)
search_button.pack(pady=20)

# Pornim bucla principală a interfeței grafice
root.mainloop()